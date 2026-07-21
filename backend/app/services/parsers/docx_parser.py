import os
import docx
from typing import List
from app.services.parsers.base_parser import BaseParser, ParsedDocument, ParsedPage


class DOCXParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found at {file_path}")

        full_text_list = []
        doc = docx.Document(file_path)

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text_list.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text_list.append(" | ".join(row_text))

        full_text = "\n\n".join(full_text_list)
        
        # Word documents are parsed as single-page structures
        pages = [ParsedPage(page_number=1, text=full_text)]
        metadata = {
            "page_count": 1,
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
        }

        return ParsedDocument(text=full_text, pages=pages, metadata=metadata)
